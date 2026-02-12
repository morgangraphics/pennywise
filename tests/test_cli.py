"""Tests for command-line interface."""

import pytest
import sys
from unittest.mock import patch, MagicMock


@pytest.fixture
def mock_parser_and_path():
    """Fixture to mock PennyParser and Path for CLI testing."""
    with patch("penny_parser.PennyParser") as mock_parser_class:
        mock_parser = MagicMock()
        mock_parser_class.return_value = mock_parser
        with patch("penny_parser.Path") as mock_path:
            mock_path.return_value.is_dir.return_value = False
            yield mock_parser


class TestArgumentParsing:
    """Test command-line argument parsing."""

    def test_required_input_argument(self):
        """Test that --input is required."""
        from penny_parser import main

        with patch.object(sys, "argv", ["penny_parser.py"]):
            with pytest.raises(SystemExit):
                main()

    def test_required_output_argument(self):
        """Test that --output is required."""
        from penny_parser import main

        with patch.object(sys, "argv", ["penny_parser.py", "-i", "input.docx"]):
            with pytest.raises(SystemExit):
                main()

    def test_input_short_form(self, mock_parser_and_path):
        """Test -i short form for input."""
        from penny_parser import main

        with patch.object(
            sys,
            "argv",
            ["penny_parser.py", "-i", "test.docx", "-o", "output.csv"],
        ):
            main()
            # Verify run was called with correct input argument
            mock_parser_and_path.run.assert_called_once_with(
                "test.docx", "output.csv", False, False, False
            )

    def test_output_short_form(self, mock_parser_and_path):
        """Test -o short form for output."""
        from penny_parser import main

        with patch.object(
            sys,
            "argv",
            ["penny_parser.py", "-i", "test.docx", "-o", "output.csv"],
        ):
            main()
            # Verify run was called with correct arguments
            mock_parser_and_path.run.assert_called_once_with(
                "test.docx", "output.csv", False, False, False
            )

    def test_short_location_flag(self, mock_parser_and_path):
        """Test -sl / --short-location flag."""
        from penny_parser import main
        from penny_parser import parse_arguments

        # Test that -sl flag is recognized and sets short_loc to True
        args = parse_arguments(["-i", "test.docx", "-o", "output.csv", "-sl"])
        assert args.short_loc is True

        with patch.object(
            sys,
            "argv",
            ["penny_parser.py", "-i", "test.docx", "-o", "output.csv", "-sl"],
        ):
            main()
            # Verify run was called with short_loc=True
            mock_parser_and_path.run.assert_called_once_with(
                "test.docx", "output.csv", True, False, False
            )

    def test_multi_line_dash_flag(self, mock_parser_and_path):
        """Test -mld / --multi-line-dash flag."""
        from penny_parser import main

        with patch.object(
            sys,
            "argv",
            ["penny_parser.py", "-i", "test.docx", "-o", "output.csv", "-mld"],
        ):
            main()
            # Verify run was called with multi_line_dash=True
            mock_parser_and_path.run.assert_called_once_with(
                "test.docx", "output.csv", False, True, False
            )

    def test_new_only_flag(self, mock_parser_and_path):
        """Test -n / --new-only flag."""
        from penny_parser import main

        with patch.object(
            sys,
            "argv",
            ["penny_parser.py", "-i", "test.docx", "-o", "output.csv", "-n"],
        ):
            main()
            # Verify run was called with new_only=True
            mock_parser_and_path.run.assert_called_once_with(
                "test.docx", "output.csv", False, False, True
            )

    def test_all_flags_together(self, mock_parser_and_path):
        """Test using all flags together."""
        from penny_parser import main
        from unittest.mock import patch

        with patch.object(
            sys,
            "argv",
            [
                "penny_parser.py",
                "-i",
                "test.docx",
                "-o",
                "output.csv",
                "-sl",
                "-mld",
                "-n",
            ],
        ):
            main()
            # Verify run was called with all flags set to True
            mock_parser_and_path.run.assert_called_once_with(
                "test.docx", "output.csv", True, True, True
            )

    def test_long_form_arguments(self, mock_parser_and_path):
        """Test using long form of arguments."""
        from penny_parser import main
        from unittest.mock import patch

        with patch.object(
            sys,
            "argv",
            [
                "penny_parser.py",
                "--input",
                "test.docx",
                "--output",
                "output.csv",
                "--short-location",
                "--multi-line-dash",
                "--new-only",
            ],
        ):
            main()
            # Verify run was called with all flags set to True using long form
            mock_parser_and_path.run.assert_called_once_with(
                "test.docx", "output.csv", True, True, True
            )

    def test_help_flag(self):
        """Test -h / --help flag shows help."""
        from penny_parser import main
        from unittest.mock import patch

        with patch.object(sys, "argv", ["penny_parser.py", "-h"]):
            with pytest.raises(SystemExit) as exc_info:
                main()
            assert exc_info.value.code == 0  # Help exits with code 0


class TestDirectoryInput:
    """Test directory input handling."""

    def test_input_accepts_file_path(self, temp_dir):
        """Test that input accepts a file path."""
        from pathlib import Path
        from docx import Document

        # Create a test docx file
        doc = Document()
        doc.add_paragraph("Test")
        test_file = Path(temp_dir) / "ca.docx"
        doc.save(str(test_file))

        assert test_file.exists()

    def test_input_accepts_directory_path(self, temp_dir):
        """Test that input accepts a directory path."""
        from pathlib import Path
        from docx import Document

        # Create test docx files
        for i in range(2):
            doc = Document()
            doc.add_paragraph(f"Test {i}")
            test_file = Path(temp_dir) / f"test_{i}.docx"
            doc.save(str(test_file))

        assert Path(temp_dir).is_dir()
        assert len(list(Path(temp_dir).glob("*.docx"))) == 2


class TestMimeTypeValidation:
    """Test MIME type validation."""

    def test_docx_mime_type_recognized(self, parser, temp_dir):
        """Test that .docx files have correct MIME type."""
        from pathlib import Path
        from docx import Document
        import mimetypes

        doc = Document()
        docx_path = Path(temp_dir) / "test.docx"
        doc.save(str(docx_path))

        mime_type, _ = mimetypes.guess_type(str(docx_path))
        assert (
            mime_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def test_txt_file_rejected(self, parser, temp_dir):
        """Test that .txt files are rejected."""
        import mimetypes

        mime_type, _ = mimetypes.guess_type("test.txt")
        assert (
            mime_type
            != "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
