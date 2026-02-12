"""Tests for PennyParser class - text parsing and sanitization."""

import pytest
from pathlib import Path


class TestSplitAndStrip:
    """Test the split_and_strip method."""

    def test_split_and_strip_default_delimiter(self, parser):
        """Test splitting with default dash delimiter."""
        result = parser.split_and_strip("Downtown Disney - World of Disney")
        assert result == ["Downtown Disney", "World of Disney"]

    def test_split_and_strip_en_dash(self, parser):
        """Test splitting with en-dash (–)."""
        result = parser.split_and_strip("Downtown Disney – World of Disney")
        assert result == ["Downtown Disney", "World of Disney"]

    def test_split_and_strip_multiple_dashes(self, parser):
        """Test splitting with multiple dashes."""
        result = parser.split_and_strip("Part One - Part Two - Part Three")
        assert len(result) == 3
        assert result == ["Part One", "Part Two", "Part Three"]

    def test_split_and_strip_no_delimiter(self, parser):
        """Test splitting text with no delimiter."""
        result = parser.split_and_strip("Single Part Only")
        assert result == ["Single Part Only"]

    def test_split_and_strip_custom_delimiter(self, parser):
        """Test splitting with custom delimiter."""
        result = parser.split_and_strip("Part A | Part B | Part C", delimiter=r"\|")
        assert result == ["Part A", "Part B", "Part C"]

    def test_split_and_strip_whitespace_handling(self, parser):
        """Test that whitespace is properly stripped."""
        result = parser.split_and_strip("  Part One  -  Part Two  ")
        assert result == ["Part One", "Part Two"]


class TestSanitizeForCsv:
    """Test the sanitize_for_csv method."""

    def test_sanitize_en_dash(self, parser):
        """Test conversion of en-dash."""
        result = parser.sanitize_for_csv("Downtown Disney – World")
        assert "–" not in result
        assert "-" in result

    def test_sanitize_em_dash(self, parser):
        """Test conversion of em-dash."""
        result = parser.sanitize_for_csv("Downtown Disney — World")
        assert "—" not in result
        assert "-" in result

    def test_sanitize_smart_quotes(self, parser):
        """Test conversion of smart quotes."""
        # Using unicode escape sequences to avoid syntax errors
        result = parser.sanitize_for_csv('It\u2019s \u201cMusée Mécanique\u201d')
        assert '\u201c' not in result
        assert '\u201d' not in result
        assert '"' in result

    def test_sanitize_trademark_copyright(self, parser):
        """Test conversion of trademark and copyright symbols."""
        result = parser.sanitize_for_csv("Disney® Corporation©")
        assert "®" not in result
        assert "©" not in result
        assert "(R)" in result
        assert "(C)" in result

    def test_sanitize_ellipsis(self, parser):
        """Test conversion of ellipsis."""
        result = parser.sanitize_for_csv("Three Dots…")
        assert "…" not in result
        assert "..." in result

    def test_sanitize_multiple_unicode(self, parser):
        """Test sanitizing multiple unicode characters at once."""
        result = parser.sanitize_for_csv('It\u2019s \u201cHello\u201d\u2013World\u2122')
        assert "'" in result
        assert '"' in result
        assert "-" in result
        assert "(TM)" in result

    def test_sanitize_already_ascii(self, parser):
        """Test that ASCII text remains unchanged."""
        text = "Normal ASCII Text"
        result = parser.sanitize_for_csv(text)
        assert result == text


class TestStripNewlines:
    """Test the strip_newlines_and_returns method."""

    def test_strip_newlines(self, parser):
        """Test removal of newline characters."""
        result = parser.strip_newlines_and_returns("Line One\nLine Two\nLine Three")
        assert "\n" not in result
        assert "Line One Line Two Line Three" == result

    def test_strip_carriage_returns(self, parser):
        """Test removal of carriage return characters."""
        result = parser.strip_newlines_and_returns("Line One\rLine Two")
        assert "\r" not in result
        assert "Line One Line Two" == result

    def test_strip_mixed_line_endings(self, parser):
        """Test removal of mixed line endings."""
        result = parser.strip_newlines_and_returns("Line One\r\nLine Two\nLine Three\r")
        assert "\n" not in result
        assert "\r" not in result
        assert " " in result  # Should have spaces between lines

    def test_strip_no_line_endings(self, parser):
        """Test text with no line endings."""
        text = "Single line text"
        result = parser.strip_newlines_and_returns(text)
        assert result == text


class TestNormalizeCellText:
    """Test the normalize_cell_text method."""

    def test_normalize_cell_text_strips_whitespace(self, parser, temp_dir):
        """Test that cell text is stripped."""
        from docx import Document
        from docx.shared import Pt

        # Create a temporary DOCX with a table cell
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = "  Cell Content  "

        docx_path = Path(temp_dir) / "test_cell.docx"
        doc.save(str(docx_path))

        # Now read and test
        doc = Document(str(docx_path))
        table = doc.tables[0]
        cell = table.rows[0].cells[0]
        result = parser.normalize_cell_text(cell)
        assert result == "Cell Content"

    def test_normalize_cell_text_empty_cell(self, parser, temp_dir):
        """Test normalizing empty cell."""
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        docx_path = Path(temp_dir) / "test_empty_cell.docx"
        doc.save(str(docx_path))

        doc = Document(str(docx_path))
        table = doc.tables[0]
        cell = table.rows[0].cells[0]
        result = parser.normalize_cell_text(cell)
        assert result == ""
