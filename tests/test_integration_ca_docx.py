"""Integration tests using the ca.docx file."""

import pytest
from pathlib import Path
import csv


class TestParsingCaDocx:
    """Integration tests parsing the actual ca.docx file."""

    @pytest.fixture
    def ca_docx_path(self):
        """Get path to ca.docx file."""
        path = Path("./pennies/labels/ca.docx")
        if not path.exists():
            pytest.skip("ca.docx file not found")
        return path

    def test_parse_ca_docx_returns_rows(self, parser, ca_docx_path):
        """Test that parsing ca.docx returns penny rows."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        assert isinstance(rows, list)
        assert len(rows) > 0

    def test_parse_ca_docx_all_rows_have_required_fields(self, parser, ca_docx_path):
        """Test that all rows have required CSV fields."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        required_fields = [
            "State",
            "City",
            "Neighborhood",
            "Location",
            "Name",
            "Orientation",
            "Type",
            "Year",
            "Position",
            "Retired",
            "Set #",
            "Quantity",
            "Need",
        ]

        for row in rows:
            for field in required_fields:
                assert field in row, f"Missing field {field} in row {row}"

    def test_parse_ca_docx_state_is_california(self, parser, ca_docx_path):
        """Test that all rows have State=California."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        for row in rows:
            assert row["State"] == "California"

    def test_parse_ca_docx_has_cities(self, parser, ca_docx_path):
        """Test that rows have city values."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        assert all(row["City"] for row in rows)

    def test_parse_ca_docx_has_locations(self, parser, ca_docx_path):
        """Test that rows have location values."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        assert all(row["Location"] for row in rows)

    def test_parse_ca_docx_has_names(self, parser, ca_docx_path):
        """Test that rows have name values."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        assert all(row["Name"] for row in rows)

    def test_parse_ca_docx_orientation_valid(self, parser, ca_docx_path):
        """Test that orientations are h or v."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        for row in rows:
            assert row["Orientation"] in [
                "h",
                "v",
            ], f"Invalid orientation: {row['Orientation']}"

    def test_parse_ca_docx_has_types(self, parser, ca_docx_path):
        """Test that all rows have type values."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        assert all(row["Type"] for row in rows)

    def test_parse_ca_docx_has_positions(self, parser, ca_docx_path):
        """Test that all rows have position values."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        assert all(row["Position"] for row in rows)

    def test_parse_ca_docx_retired_is_yes_or_empty(self, parser, ca_docx_path):
        """Test that retired is either 'Yes' or empty string."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        for row in rows:
            assert row["Retired"] in ["Yes", ""]

    def test_parse_ca_docx_no_unicode_chars(self, parser, ca_docx_path):
        """Test that unicode characters are sanitized."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        # List of problematic unicode characters that should be sanitized
        problematic_chars = [
            "\u2013",
            "\u2014",
            "\u201c",
            "\u201d",
            "\u00ae",
            "\u2122",
            "\u00a9",
        ]
        for row in rows:
            for field in ["Name", "Location", "Neighborhood", "City"]:
                for char in problematic_chars:
                    assert (
                        char not in row[field]
                    ), f"Found unsanitized char '{char}' in {field}: {row[field]}"

    def test_parse_ca_docx_no_newlines(self, parser, ca_docx_path):
        """Test that newlines are removed from fields."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        for row in rows:
            for field in ["Name", "Location", "Neighborhood", "City"]:
                assert "\n" not in row[field]
                assert "\r" not in row[field]

    def test_write_csv_output(self, parser, ca_docx_path, temp_dir):
        """Test writing parsed rows to CSV."""
        rows = parser.parse_docx_to_rows(str(ca_docx_path))
        output_path = Path(temp_dir) / "output.csv"
        parser.write_csv(rows, str(output_path))

        assert output_path.exists()
        with open(output_path, "r") as f:
            reader = csv.DictReader(f)
            csv_rows = list(reader)
            assert len(csv_rows) == len(rows)

    def test_parse_with_short_location_flag(self, parser, ca_docx_path):
        """Test parsing with short_location flag."""
        parser.short_location = False
        rows_full = parser.parse_docx_to_rows(str(ca_docx_path))

        parser.short_location = True
        rows_short = parser.parse_docx_to_rows(str(ca_docx_path))

        # Both should have same number of rows
        assert len(rows_full) == len(rows_short)
        # Short location should have shorter or equal location strings
        for full_row, short_row in zip(rows_full, rows_short):
            assert len(short_row["Location"]) <= len(full_row["Location"])

        parser.short_location = False

    def test_parse_with_multi_line_dash_flag(self, parser, ca_docx_path):
        """Test parsing with multi_line_dash flag."""
        parser.multi_line_dash = False
        rows_no_dash = parser.parse_docx_to_rows(str(ca_docx_path))

        parser.multi_line_dash = True
        rows_with_dash = parser.parse_docx_to_rows(str(ca_docx_path))

        # Both should have same number of rows
        assert len(rows_no_dash) == len(rows_with_dash)

        parser.multi_line_dash = False

    def test_parse_with_new_only_flag(self, parser, ca_docx_path):
        """Test parsing with new_only flag."""
        # First parse: should get all pennies
        parser.new_only = False
        all_rows = parser.parse_docx_to_rows(str(ca_docx_path))

        # Add some to database, normalizing keys to lowercase
        if len(all_rows) > 0:
            penny_to_add = {
                k.lower(): v for k, v in all_rows[0].items()
            }
            parser.db.add_penny(penny_to_add)

        # Second parse with new_only: should get fewer pennies
        parser.new_only = True
        new_rows = parser.parse_docx_to_rows(str(ca_docx_path))

        # new_only should return fewer rows
        assert len(new_rows) < len(all_rows)

        parser.new_only = False


class TestH2H3MapBuilder:
    """Test the build_h2_h3_map method with ca.docx."""

    def test_h2_h3_map_built(self, parser):
        """Test that H2/H3 map is built correctly."""
        ca_docx_path = Path("./pennies/labels/ca.docx")
        if not ca_docx_path.exists():
            pytest.skip("ca.docx file not found")

        from docx import Document

        doc = Document(str(ca_docx_path))
        h2_h3_map = parser.build_h2_h3_map(doc)

        # Map should have entries
        assert isinstance(h2_h3_map, dict)
        assert len(h2_h3_map) > 0

    def test_h2_h3_map_identifies_neighborhoods(self, parser):
        """Test that H2s with H3s below are identified as neighborhoods."""
        ca_docx_path = Path("./pennies/labels/ca.docx")
        if not ca_docx_path.exists():
            pytest.skip("ca.docx file not found")

        from docx import Document

        doc = Document(str(ca_docx_path))
        h2_h3_map = parser.build_h2_h3_map(doc)

        # Some H2s should have H3s (True), some shouldn't (False)
        has_h3_true = sum(1 for v in h2_h3_map.values() if v is True)
        has_h3_false = sum(1 for v in h2_h3_map.values() if v is False)

        # Should have at least some of each
        assert has_h3_true > 0 and has_h3_false > 0


class TestRowsFromVerticalPairs:
    """Test the rows_from_vertical_pairs method."""

    def test_vertical_pairs_yields_tuples(self, parser, temp_dir):
        """Test that rows_from_vertical_pairs yields tuples."""
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=4, cols=7)
        # Add some test data
        for i in range(4):
            table.cell(i, 0).text = f"Cell {i}"

        docx_path = Path(temp_dir) / "test_vertical.docx"
        doc.save(str(docx_path))

        doc = Document(str(docx_path))
        table = doc.tables[0]

        pairs = list(parser.rows_from_vertical_pairs(table))
        assert len(pairs) > 0

        for pair in pairs:
            assert len(pair) == 3  # (cell1, cell2, position)

    def test_vertical_pairs_skips_spacer_columns(self, parser, temp_dir):
        """Test that odd-indexed columns (spacers) are skipped."""
        from docx import Document

        doc = Document()
        # 7-column format: col0=data, col1=spacer, col2=data, col3=spacer, col4=data, col5=spacer, col6=data
        doc.add_table(rows=2, cols=7)

        docx_path = Path(temp_dir) / "test_spacer.docx"
        doc.save(str(docx_path))

        doc = Document(str(docx_path))
        table = doc.tables[0]

        pairs = list(parser.rows_from_vertical_pairs(table))
        # With 7 columns, should have 4 data columns (0, 2, 4, 6)
        # Each data column yields 1 pair (2 rows / 2)
        assert len(pairs) == 4
