"""Tests for PennyParser class - state parsing and retired detection."""

from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


class TestParseStateFromFilename:
    """Test the parse_state_from_filename method."""

    def test_parse_california(self, parser):
        """Test parsing California state."""
        result = parser.parse_state_from_filename("ca.docx")
        assert result == "California"

    def test_parse_colorado(self, parser):
        """Test parsing Colorado state."""
        result = parser.parse_state_from_filename("co.docx")
        assert result == "Colorado"

    def test_parse_new_york(self, parser):
        """Test parsing New York state."""
        result = parser.parse_state_from_filename("ny.docx")
        assert result == "New York"

    def test_parse_uppercase(self, parser):
        """Test that uppercase filenames are handled."""
        result = parser.parse_state_from_filename("CA.docx")
        assert result == "California"

    def test_parse_mixed_case(self, parser):
        """Test mixed case filenames."""
        result = parser.parse_state_from_filename("Ca.docx")
        assert result == "California"

    def test_parse_path_with_directory(self, parser):
        """Test parsing filename from full path."""
        result = parser.parse_state_from_filename("/path/to/ca.docx")
        assert result == "California"

    def test_parse_windows_path(self, parser):
        """Test parsing filename from Windows path."""
        result = parser.parse_state_from_filename("C:/Users/Documents/ca.docx")
        assert result == "California"

    def test_parse_unknown_state(self, parser):
        """Test parsing unknown state returns capitalized version."""
        result = parser.parse_state_from_filename("xx.docx")
        assert result == "Xx"

    def test_all_us_states(self, parser):
        """Test that all US states in state_map are parsed correctly."""
        states = [
            ("al", "Alabama"),
            ("ak", "Alaska"),
            ("az", "Arizona"),
            ("ar", "Arkansas"),
            ("ca", "California"),
            ("co", "Colorado"),
            ("ct", "Connecticut"),
            ("dc", "District of Columbia"),
            ("fl", "Florida"),
            ("ga", "Georgia"),
            ("hi", "Hawaii"),
            ("il", "Illinois"),
            ("ny", "New York"),
            ("tx", "Texas"),
            ("wa", "Washington"),
            ("wy", "Wyoming"),
        ]
        for abbrev, full_name in states:
            result = parser.parse_state_from_filename(f"{abbrev}.docx")
            assert result == full_name, f"Failed for {abbrev}"

    def test_parse_abbreviation_with_dash_suffix(self, parser):
        """Test parsing state abbreviation with dash suffix."""
        result = parser.parse_state_from_filename("ca-new.docx")
        assert result == "California"

    def test_parse_abbreviation_with_underscore_suffix(self, parser):
        """Test parsing state abbreviation with underscore suffix."""
        result = parser.parse_state_from_filename("ca_backup.docx")
        assert result == "California"

    def test_parse_abbreviation_with_dot_suffix(self, parser):
        """Test parsing state abbreviation with dot suffix."""
        result = parser.parse_state_from_filename("co.old.docx")
        assert result == "Colorado"

    def test_parse_abbreviation_with_multiple_suffixes(self, parser):
        """Test parsing state abbreviation with multiple suffix words."""
        result = parser.parse_state_from_filename("ny-backup-2024.docx")
        assert result == "New York"

    def test_parse_full_state_name_single_word(self, parser):
        """Test parsing full state name for single-word states."""
        result = parser.parse_state_from_filename("massachusetts.docx")
        assert result == "Massachusetts"

    def test_parse_full_state_name_with_suffix(self, parser):
        """Test parsing full state name with suffix."""
        result = parser.parse_state_from_filename("massachusetts-old.docx")
        assert result == "Massachusetts"

    def test_parse_full_state_california(self, parser):
        """Test parsing full California state name."""
        result = parser.parse_state_from_filename("california.docx")
        assert result == "California"

    def test_parse_full_state_california_with_suffix(self, parser):
        """Test parsing full California state name with suffix."""
        result = parser.parse_state_from_filename("california-new.docx")
        assert result == "California"

    def test_parse_multiword_state_with_dots(self, parser):
        """Test parsing multi-word state name with dots."""
        result = parser.parse_state_from_filename("new.york.docx")
        assert result == "New York"

    def test_parse_multiword_state_with_dashes(self, parser):
        """Test parsing multi-word state name with dashes."""
        result = parser.parse_state_from_filename("new-york.docx")
        assert result == "New York"

    def test_parse_multiword_state_new_mexico_dots(self, parser):
        """Test parsing New Mexico with dots."""
        result = parser.parse_state_from_filename("new.mexico.docx")
        assert result == "New Mexico"

    def test_parse_multiword_state_new_mexico_dashes(self, parser):
        """Test parsing New Mexico with dashes."""
        result = parser.parse_state_from_filename("new-mexico.docx")
        assert result == "New Mexico"

    def test_parse_multiword_state_north_carolina_dots(self, parser):
        """Test parsing North Carolina with dots."""
        result = parser.parse_state_from_filename("north.carolina.docx")
        assert result == "North Carolina"

    def test_parse_multiword_state_north_carolina_dashes(self, parser):
        """Test parsing North Carolina with dashes."""
        result = parser.parse_state_from_filename("north-carolina.docx")
        assert result == "North Carolina"

    def test_parse_multiword_state_with_suffix(self, parser):
        """Test parsing multi-word state name with suffix."""
        result = parser.parse_state_from_filename("new-york-backup.docx")
        assert result == "New York"

    def test_collision_in_not_abbreviation(self, parser):
        """Test that 'in' in 'indiana' doesn't match Indiana abbreviation."""
        result = parser.parse_state_from_filename("indiana.docx")
        assert result == "Indiana"

    def test_collision_or_word_vs_oregon(self, parser):
        """Test that standalone 'or' word doesn't incorrectly match when not state."""
        # 'or' is Oregon abbreviation, but 'oregon' should match as full state name
        result = parser.parse_state_from_filename("oregon.docx")
        assert result == "Oregon"

    def test_false_positive_common_word_in_filename(self, parser):
        """Test that common words that might be state abbreviations don't match."""
        # 'me' is Maine abbreviation, but 'readme' should not match
        result = parser.parse_state_from_filename("readme.docx")
        assert result == "Readme"  # Fallback capitalization

    def test_false_positive_partial_state_name(self, parser):
        """Test that partial state names don't incorrectly match."""
        # 'wash' should not match Washington
        result = parser.parse_state_from_filename("wash.docx")
        assert result == "Wash"  # Fallback capitalization


class TestCellIsRetired:
    """Test the cell_is_retired method."""

    def test_retired_cell_f2dbdb_color(self, parser, temp_dir):
        """Test detection of retired cell with f2dbdb color."""
        from docx import Document
        from docx.oxml import OxmlElement

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = "Test"

        # Set the cell shading to f2dbdb
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "f2dbdb")
        tcPr.append(shd)

        docx_path = Path(temp_dir) / "test_retired.docx"
        doc.save(str(docx_path))

        doc = Document(str(docx_path))
        cell = doc.tables[0].rows[0].cells[0]
        assert parser.cell_is_retired(cell) is True

    def test_active_cell_no_shading(self, parser, temp_dir):
        """Test that active cell without shading is not retired."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = "Test"

        docx_path = Path(temp_dir) / "test_active.docx"
        doc.save(str(docx_path))

        doc = Document(str(docx_path))
        cell = doc.tables[0].rows[0].cells[0]
        assert parser.cell_is_retired(cell) is False

    def test_active_cell_different_color(self, parser, temp_dir):
        """Test that active cell with non-retired color is not retired."""
        from docx import Document
        from docx.oxml import OxmlElement

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = "Test"

        # Set a different color
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "FFFFFF")  # White
        tcPr.append(shd)

        docx_path = Path(temp_dir) / "test_white.docx"
        doc.save(str(docx_path))

        doc = Document(str(docx_path))
        cell = doc.tables[0].rows[0].cells[0]
        assert parser.cell_is_retired(cell) is False

    def test_retired_uppercase_hex_color(self, parser, temp_dir):
        """Test that uppercase hex colors are recognized."""
        from docx import Document
        from docx.oxml import OxmlElement

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = "Test"

        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F2DBDB")  # Uppercase
        tcPr.append(shd)

        docx_path = Path(temp_dir) / "test_retired_uppercase.docx"
        doc.save(str(docx_path))

        doc = Document(str(docx_path))
        cell = doc.tables[0].rows[0].cells[0]
        assert parser.cell_is_retired(cell) is True
