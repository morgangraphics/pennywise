#!/usr/bin/env python3
"""
penny_parser.py

CLI tool for parsing specially formatted Word (.docx) documents containing
vertical-stacked label tables into normalized CSV files.

Usage:
    python penny_parser.py --input il.docx --output il.csv
"""

import csv
import argparse
import mimetypes
import re
import logging
import sys
import io
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from penny_database import PennyDatabase


class PennyParser:
    """Parser for pressed-penny DOCX documents."""

    # Separator dash pattern (avoid splitting in-word like "two-way" or "Buc-ees")
    dash_regex = r"(?<!\w)[-–—](?!\w)"

    def __init__(self, log_file: str = "penny_parser.log", db_file: str = "pennies.db"):
        """
        Initialize the parser with logging configuration.

        Args:
            log_file (str): Path to the log file.
            db_file (str): Path to SQLite database file.
        """
        self.log_file = log_file
        self.logger = self.setup_logging(log_file, with_console=True)
        self.labels_logger = self.setup_logging(
            "labels.log", logger_name=f"{__name__}.labels", with_console=False
        )
        self.short_location = False
        self.write_mode = "w"
        self.new_only = False
        self.multi_line_dash = False
        self.db = PennyDatabase(db_file)
        self.continuation_words = (
            "And",
            "and",
            "&",
            "Of",
            "of",
        )  # Words/symbols that indicate continuation without separator
        self.state_map = {
            "al": "Alabama",
            "ak": "Alaska",
            "az": "Arizona",
            "ar": "Arkansas",
            "ca": "California",
            "co": "Colorado",
            "ct": "Connecticut",
            "de": "Delaware",
            "dc": "District of Columbia",
            "fl": "Florida",
            "ga": "Georgia",
            "hi": "Hawaii",
            "id": "Idaho",
            "il": "Illinois",
            "in": "Indiana",
            "ia": "Iowa",
            "ks": "Kansas",
            "ky": "Kentucky",
            "la": "Louisiana",
            "me": "Maine",
            "md": "Maryland",
            "ma": "Massachusetts",
            "mi": "Michigan",
            "mn": "Minnesota",
            "ms": "Mississippi",
            "mo": "Missouri",
            "mt": "Montana",
            "ne": "Nebraska",
            "nv": "Nevada",
            "nh": "New Hampshire",
            "nj": "New Jersey",
            "nm": "New Mexico",
            "ny": "New York",
            "nc": "North Carolina",
            "nd": "North Dakota",
            "oh": "Ohio",
            "ok": "Oklahoma",
            "or": "Oregon",
            "pa": "Pennsylvania",
            "ri": "Rhode Island",
            "sc": "South Carolina",
            "sd": "South Dakota",
            "tn": "Tennessee",
            "tx": "Texas",
            "ut": "Utah",
            "vt": "Vermont",
            "va": "Virginia",
            "wa": "Washington",
            "wv": "West Virginia",
            "wi": "Wisconsin",
            "wy": "Wyoming",
            # Add territories here if desired
        }

    def setup_logging(
        self, log_file: str, logger_name: str = None, with_console: bool = True
    ):
        """
        Configure logging to output to file and optionally console.

        Args:
            log_file (str): Path to the log file.
            logger_name (str): Name for the logger. Defaults to __name__.
            with_console (bool): Whether to include console output. Defaults to True.
        """
        logger = logging.getLogger(logger_name or __name__)
        logger.setLevel(logging.DEBUG)

        # Remove existing handlers
        logger.handlers = []

        # Console handler - INFO and above (if enabled)
        if with_console:
            console_stream = sys.stderr
            try:
                console_stream = io.TextIOWrapper(
                    sys.stderr.buffer, encoding="utf-8", errors="replace"
                )
            except Exception:
                console_stream = sys.stderr
            console_handler = logging.StreamHandler(console_stream)
            console_handler.setLevel(logging.INFO)
            console_formatter = logging.Formatter("%(levelname)s: %(message)s")
            console_handler.setFormatter(console_formatter)
            logger.addHandler(console_handler)

        # File handler - DEBUG and above (more verbose)
        file_handler = logging.FileHandler(log_file, encoding="utf-8")
        file_handler.setLevel(logging.DEBUG)
        file_formatter = logging.Formatter(
            "%(asctime)s - %(levelname)s - %(funcName)s:%(lineno)d - %(message)s"
        )
        file_handler.setFormatter(file_formatter)
        logger.addHandler(file_handler)

        # Prevent propagation to parent logger if this is a child logger
        if logger_name and logger_name != __name__:
            logger.propagate = False

        return logger

    def parse_state_from_filename(self, filename: str) -> str:
        """
        Extract state name from filename using state_map as source of truth.
        
        Handles various filename formats:
        - ca.docx, ca-new.docx, ca_backup.docx → California
        - massachusetts.docx, massachusetts-old.docx → Massachusetts
        - new.york.docx, new-york.docx → New York

        Args:
            filename (str): The DOCX filename.

        Returns:
            str: Full state name, e.g., 'Colorado'.
        """
        stem = Path(filename).stem.lower()
        # Normalize: replace separators with spaces and also keep original
        normalized_stem = stem.replace('-', ' ').replace('_', ' ').replace('.', ' ')
        
        # First, check if any full state name (value) appears in normalized stem
        # Sort by length descending to match longer names first (e.g., "new york" before "ne")
        sorted_states = sorted(self.state_map.items(), key=lambda x: len(x[1]), reverse=True)
        for abbrev, state_name in sorted_states:
            state_lower = state_name.lower()
            if state_lower in normalized_stem or state_lower.replace(' ', '') in stem:
                return state_name
        
        # Second, check if any state abbreviation (key) appears in the stem
        # Sort by length descending to prioritize longer abbreviations
        sorted_abbrevs = sorted(self.state_map.items(), key=lambda x: len(x[0]), reverse=True)
        for abbrev, state_name in sorted_abbrevs:
            # Match abbreviation as a complete token (word boundary)
            if stem == abbrev or stem.startswith(abbrev + '-') or stem.startswith(abbrev + '_') or stem.startswith(abbrev + '.'):
                return state_name
        
        # Fallback: capitalize the stem
        return stem.capitalize()

    def split_and_strip(self, text: str, delimiter: str = None) -> list:
        """
        Split text on a delimiter and strip whitespace from each resulting part.

        Args:
            text (str): The text to split.
            delimiter (str): Regex pattern to split on. Defaults to self.dash_regex.

        Returns:
            list: List of stripped text parts.
        """
        if delimiter is None:
            delimiter = self.dash_regex
        return [part.strip() for part in re.split(delimiter, text)]

    def cell_is_retired(self, cell) -> bool:
        """
        Determine whether a DOCX table cell is 'retired' by checking whether the
        shading fill color is pink/rose.

        Args:
            cell (docx.table._Cell): A DOCX table cell.

        Returns:
            bool: True if the cell background corresponds to a retired label.
        """
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is not None:
            fill = shd.get(qn("w:fill"))
            if fill:
                fill = fill.lower()
                # Colors typically used in the documents for retired pennies
                # return fill in ("ffcccc", "ffc0cb", "f4cccc", "e6b8af")
                return fill in ("f2dbdb")
        return False

    def is_table_of_contents(self, block) -> bool:
        """
        Determine if a block is part of a table of contents.

        Checks for paragraph style names containing 'TOC' (case-insensitive).
        Actual Heading paragraphs use "Heading" styles, not "TOC" styles.

        Args:
            block: A paragraph or table object.

        Returns:
            bool: True if the block is part of TOC.
        """
        if (
            hasattr(block, "style")
            and block.style
            and "toc" in block.style.name.lower()
        ):
            return True

        return False

    def build_h2_h3_map(self, document) -> dict:
        """
        Build a map of H2 headings and whether they have H3s below them.

        This allows us to determine if an H2 is just a neighborhood (has H3s below)
        or both neighborhood AND location (no H3s below).

        Args:
            document: A python-docx Document object.

        Returns:
            dict: Mapping of H2 text -> bool (True if H2 has H3s below it)
        """
        h2_h3_map = {}
        current_h2 = None

        for para in document.paragraphs:
            if not hasattr(para, "style") or not para.style:
                continue

            style_name = para.style.name.lower()
            text = para.text.strip()

            if not text:
                continue

            if "heading 2" in style_name:
                current_h2 = text
                if current_h2 not in h2_h3_map:
                    h2_h3_map[current_h2] = False
            elif "heading 3" in style_name and current_h2:
                h2_h3_map[current_h2] = True

        return h2_h3_map

    def sanitize_for_csv(self, text: str) -> str:
        """
        Replace problematic unicode characters with ASCII equivalents for safe CSV output.

        Args:
            text (str): The text to sanitize.

        Returns:
            str: Sanitized text safe for CSV.
        """
        # Common unicode replacements
        replacements = {
            "\u2013": "-",  # en dash (U+2013)
            "\u2014": "-",  # em dash (U+2014)
            "\u2019": "'",  # right single quotation mark / apostrophe (U+2019)
            "\u2018": "'",  # left single quotation mark (U+2018)
            "\u201c": '"',  # left double quotation mark (U+201C)
            "\u201d": '"',  # right double quotation mark (U+201D)
            "\u02bc": "'",  # modifier letter apostrophe (U+02BC)
            "\u00b4": "'",  # acute accent (U+00B4)
            "\u201b": "'",  # reversed comma quotation mark (U+201B)
            "\u2032": "'",  # prime symbol (U+2032)
            "\u2026": "...",  # ellipsis (U+2026)
            "\u00ae": "(R)",  # registered trademark (U+00AE)
            "\u2122": "(TM)",  # trademark (U+2122)
            "\u00a9": "(C)",  # copyright (U+00A9)
            "\u00b0": "deg",  # degree symbol (U+00B0)
            "&": "&",  # ampersand (already safe, but explicit)
        }

        result = text
        for unicode_char, ascii_char in replacements.items():
            result = result.replace(unicode_char, ascii_char)

        return result

    def normalize_cell_text(self, cell) -> str:
        """
        Extract text from a DOCX table cell and strip whitespace.

        Args:
            cell (docx.table._Cell): Table cell.

        Returns:
            str: Cleaned text string.
        """
        return cell.text.strip()

    def strip_newlines_and_returns(self, text: str) -> str:
        """
        Remove newline and carriage return characters from text.

        Args:
            text (str): The text to clean.

        Returns:
            str: Text with \\n and \\r characters removed.
        """
        return text.replace("\n", " ").replace("\r", " ")

    def detect_city_location_neighborhood(self, label_text: str):
        """
        Parse a top label cell text into (city, location, neighborhood).

        The label text is multi-line with structure:
        Line 1: City
        Line 2: Neighborhood (or Neighborhood - Location if has dash)
        Line 3+: Location (may continue with dashes or spaces)

        EXAMPLE:
            Santa Monica
            Santa Monica Pier
            Bubba Gump Shrimp Company
            =>
            ('Santa Monica', 'Bubba Gump Shrimp Company', 'Santa Monica Pier')

        Args:
            label_text (str): The raw label text from table cell.

        Returns:
            tuple: (city, location, neighborhood)
        """
        city = ""
        location = ""
        neighborhood = ""

        lines = label_text.splitlines()

        for i, line in enumerate(lines):
            if i == 0:
                # First line is always the city
                city = self.sanitize_for_csv(line.strip())

            elif i == 1 and len(lines) == 2:
                # Two lines total: City and Location (could have dash for neighborhood)
                if re.search(self.dash_regex, line):
                    # Has dash: Line is "Neighborhood – Location"
                    temp = self.split_and_strip(line)
                    neighborhood = self.sanitize_for_csv(temp[0])
                    if self.short_location:
                        # Keep only the short location (after dash)
                        location = self.sanitize_for_csv(" ".join(temp[1:]).strip())
                    else:
                        # Keep full "Neighborhood - Location" format
                        location = self.sanitize_for_csv(" - ".join(temp).strip())
                else:
                    # No dash: Line is just Location
                    location = self.sanitize_for_csv(line.strip())

            elif i == 1 and len(lines) > 2:
                # More than two lines: Second line is neighborhood
                if re.search(self.dash_regex, line):
                    # Has dash: Split into neighborhood and location part
                    temp = self.split_and_strip(line)
                    neighborhood = self.sanitize_for_csv(temp[0])
                    if self.short_location:
                        # Keep only the short location (after dash)
                        location = self.sanitize_for_csv(" ".join(temp[1:]).strip())
                    else:
                        # Keep full "Neighborhood - Location" format
                        location = self.sanitize_for_csv(" - ".join(temp).strip())
                else:
                    # No dash: Line is neighborhood only, location comes from line 3+
                    neighborhood = self.sanitize_for_csv(line.strip())
                    location = ""  # Will be filled by subsequent lines

            elif i >= 2:
                # Third line onwards: Additional location details
                if re.search(self.dash_regex, line):
                    # Has dash in line
                    temp = self.split_and_strip(line)
                    if neighborhood == "":
                        neighborhood = self.sanitize_for_csv(temp[0])

                    if self.short_location:
                        # Keep only the location part (after dash)
                        location = (
                            self.sanitize_for_csv(" ".join(temp[1:]).strip())
                            if len(temp) > 1
                            else self.sanitize_for_csv(line.strip())
                        )
                    else:
                        # Append full line to location
                        if location != "":
                            location = (
                                location + " - " + self.sanitize_for_csv(line.strip())
                            )
                        else:
                            location = self.sanitize_for_csv(line.strip())
                else:
                    # No dash: Continuation of location
                    # Check if line starts with continuation words (And, Of, etc.)
                    spacer = (
                        " "
                        if line.strip().startswith(self.continuation_words)
                        else " - "
                    )

                    if self.short_location:
                        # Use this line as the location (last line wins)
                        location = self.sanitize_for_csv(line).strip()
                    else:
                        # Prepend neighborhood if location is empty
                        if location == "" and neighborhood != "":
                            location = (
                                neighborhood
                                + spacer
                                + self.sanitize_for_csv(line).strip()
                            )
                        elif location != "":
                            # Append to existing location
                            location = (
                                location + spacer + self.sanitize_for_csv(line).strip()
                            )
                        else:
                            location = self.sanitize_for_csv(line).strip()

        return city, location, neighborhood

    def detect_orientation_and_type(self, label_text: str):
        """
        Parse a label string into (orientation, type, name).

        EXAMPLE:
            'Denver Colorado - Mile High City (h) Copper Penny'
            =>
            ('h', 'Copper Penny', 'Denver Colorado - Mile High City')

        Args:
            label_text (str): The raw label text.

        Returns:
            tuple: (orientation, type_str, name_str)
        """
        orient_match = re.search(r"\((h|v)\)", label_text, re.IGNORECASE)
        if not orient_match:
            return "", "", label_text

        orientation = orient_match.group(1).lower()
        lines = label_text.splitlines()

        if len(lines) == 1:
            name_str = label_text[: orient_match.start()].strip()

        # Handles mutli-line names
        if len(lines) > 1:
            if self.multi_line_dash:
                # Allow dash in multi-line names
                name_str = (
                    self.strip_newlines_and_returns(
                        lines[0] + " - " + " ".join(lines[1 : len(lines) - 1])
                    )
                    if len(lines) > 1
                    else ""
                )
            else:
                name_str = (
                    self.strip_newlines_and_returns(" ".join(lines[0 : len(lines) - 1]))
                    if len(lines) > 1
                    else ""
                )

        type_str = label_text[orient_match.end() :].strip()

        return orientation, type_str, name_str

    def rows_from_vertical_pairs(self, table):
        """
        Yield logical rows from a DOCX table by grouping vertically stacked cells.

        Each column in the table contains vertically paired entries:
        (row0,row1), (row2,row3), etc.

        Skips even-numbered columns (2, 4, 6) which are spacer columns in the 7-column format.
        Only processes odd-numbered columns (1, 3, 5, 7) which contain actual label data.

        Position index resets for each pair of rows:
        - Rows 0-1: positions 1-4
        - Rows 2-3: positions 5-8
        - Rows 4-5: positions 9-12, etc.

        Args:
            table (docx.table.Table): The DOCX table.

        Yields:
            tuple: (cell1, cell2, position_index)
        """
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        # Count only non-spacer columns (even col_index values)
        num_data_cols = (num_cols + 1) // 2  # For 7 columns: 4 data columns

        def safe_cell(row_idx: int, col_idx: int):
            try:
                return table.cell(row_idx, col_idx)
            except IndexError:
                return None

        data_col_num = 0  # Track which data column we're on (0, 1, 2, 3...)
        for col_index in range(num_cols):
            # Skip even-numbered columns (0-indexed, so even col_index = even column number)
            if col_index % 2 == 0:  # Only process odd columns (1, 3, 5, 7...)
                for r in range(0, num_rows, 2):
                    cell1 = safe_cell(r, col_index)
                    cell2 = safe_cell(r + 1, col_index) if r + 1 < num_rows else None
                    if cell1 is None:
                        continue
                    # Position calculation: data_col_num + 1 + (row_pair_index * num_data_cols)
                    row_pair_index = r // 2
                    position = data_col_num + 1 + (row_pair_index * num_data_cols)
                    yield (cell1, cell2, position)
                data_col_num += 1

    def parse_docx_to_rows(self, filename: str) -> list:
        """
        Parse a DOCX file of pressed-penny label tables into CSV row dictionaries.

        This function implements all custom rules:
          - State is derived from filename
          - City from paragraph headers
          - Neighborhood left empty by default
          - Location from table header
          - Position determined by *column index* counting empty columns
          - Retired detected by pink/rose cell shading
          - Orientation & Type parsed from label
          - Quantity defaults to 1

        Args:
            filename (str): The .docx input file.

        Returns:
            list of dict: Parsed CSV rows.
        """
        document = Document(filename)
        state = self.parse_state_from_filename(filename)

        # Build map of H2s and whether they have H3s below them
        h2_has_h3_map = self.build_h2_h3_map(document)

        csv_rows = []
        row_dict = {
            "State": "",
            "City": "",
            "Neighborhood": "",
            "Location": "",
            "Orientation": "",
            "Name": "",
            "Type": "",
            "Year": "",
            "Position": "",
            "Retired": "",
            "Set #": "",
            "Quantity": 1,
            "Need": "",
        }

        current_set = 0  # Counter for sets within a year
        last_year = None  # Track the last year to detect when it changes

        # Initialize row_data so tables before any Heading 1 don't crash
        row_data = row_dict.copy()

        # Iterate through document body elements in order (maintains document flow)
        for element in document.element.body:
            # Check if it's a paragraph
            if element.tag.endswith("p"):
                # Get the paragraph object
                para = None
                for p in document.paragraphs:
                    if p._element == element:
                        para = p
                        break

                if para is None:
                    continue

                # Skip TOC blocks
                if self.is_table_of_contents(para):
                    continue

                text = para.text.strip()
                if not text:
                    continue

                # Determine heading level and update appropriate variable from Document Map Headings
                if hasattr(para, "style") and para.style:
                    style_name = para.style.name.lower()

                    if "heading 1" in style_name:
                        # City
                        row_data = row_dict.copy()
                        row_data.update(
                            {
                                "State": self.sanitize_for_csv(state),
                                "City": self.sanitize_for_csv(text),
                            }
                        )
                        current_set = 0  # Reset set counter for new city
                        last_year = None  # reset year

                    elif "heading 2" in style_name:
                        # H2 can be a neighborhood OR a neighborhood+location

                        self.logger.info(f"LEVEL 2 Heading detected: {text}")

                        # Check if this H2 has H3s below it
                        h2_has_h3 = h2_has_h3_map.get(text, False)

                        if h2_has_h3:
                            # H2 is just the neighborhood, H3s will be the locations
                            row_data.update(
                                {
                                    #"Neighborhood": self.sanitize_for_csv(text),
                                     "Neighborhood": "",
                                    "Location": "",
                                }
                            )
                        else:
                            # H2 is the location (Neighborhood intentionally empty)
                            if self.short_location:
                                # short_location: Location is empty, only Neighborhood is set
                                row_data.update(
                                    {
                                        #"Neighborhood": self.sanitize_for_csv(text),
                                        "Neighborhood": "",
                                        "Location": "",
                                    }
                                )
                            else:
                                # full format: Location matches Neighborhood
                                row_data.update(
                                    {
                                        #"Neighborhood": self.sanitize_for_csv(text),
                                        "Neighborhood": "",
                                        "Location": self.sanitize_for_csv(text),
                                    }
                                )

                    elif "heading 3" in style_name:
                        """
                        H3 can be:
                        1. Just a location (when H2 has multiple H3s)
                           e.g. H2="Boardwalk" -> H3="Boardwalk Arcade"

                        2. Neighborhood - Location (when H3 contains dash)
                           e.g. H3="Beech Street – Alcatraz Mini Mart"
                           Parse: neighborhood="Beech Street", location="Alcatraz Mini Mart"
                        """
                        self.logger.info(f"LEVEL 3 Heading detected: {text}")

                        # Save the current neighborhood from H2 before updating location
                        current_neighborhood = row_data.get("Neighborhood", "")

                        # Check if H3 contains a dash (neighborhood - location format)
                        if re.search(self.dash_regex, text):
                            # Split on dash
                            parts = self.split_and_strip(text)
                            neighborhood = parts[0]
                            location = " ".join(parts[1:])
                        else:
                            # H3 is just the location, use current neighborhood from H2
                            #neighborhood = current_neighborhood
                            neighborhood = ""
                            location = text

                        # Update row_data with parsed neighborhood and location
                        # Preserves State, City, Year, and other context from parent headings
                        if self.short_location:
                            # short_location: Only use the location part after the dash
                            row_data.update(
                                {
                                    #"Neighborhood": self.sanitize_for_csv(neighborhood),
                                    "Neighborhood": "",
                                    "Location": self.sanitize_for_csv(location),
                                }
                            )
                        else:
                            # full format: Combine neighborhood and location
                            full_location = (
                                f"{neighborhood} - {location}"
                                if neighborhood
                                else location
                            )
                            row_data.update(
                                {
                                    #"Neighborhood": self.sanitize_for_csv(neighborhood),
                                    "Neighborhood": "",
                                    "Location": self.sanitize_for_csv(full_location),
                                }
                            )

                        # Reset year after reaching a new level 3 heading
                        last_year = None

                        self.logger.debug(f"Level 3: {row_data.copy()}")

                    elif "heading 4" in style_name:

                        self.logger.info(f"LEVEL 4: Heading detected: {text}")

                        # Year heading - check if it's a new year to reset set counter
                        if text != last_year:
                            current_set = 0
                            last_year = text
                        row_data.update({"Year": text})

                        self.logger.debug(f"Level 4: {row_data.copy()}")

            # Check if it's a table
            elif element.tag.endswith("tbl"):
                # Get the table object
                table = None

                # copy of row_data to avoid mutation issues
                cell_data = row_data.copy()

                # tmp storage to look for multiple instances with different types
                position_stor = {}

                for t in document.tables:
                    if t._element == element:
                        table = t
                        break

                if table is None:
                    continue

                # Increment set counter for each table under a year
                if cell_data.get("Year"):
                    current_set += 1

                # Process all vertically paired rows
                for cell1, cell2, position in self.rows_from_vertical_pairs(table):

                    txt1 = self.normalize_cell_text(cell1)
                    txt2 = self.normalize_cell_text(cell2) if cell2 else ""

                    # skip if both cells are empty
                    if not txt1 and not txt2:
                        continue

                    # if penny is marked as retired
                    retired = self.cell_is_retired(cell1) and (
                        cell2 and self.cell_is_retired(cell2)
                    )

                    # Top Labels - Parse and validate against Document Map
                    if txt1:
                        city, location, neighborhood = (
                            self.detect_city_location_neighborhood(txt1)
                        )

                        # Human-readable logging of parsed vs document map
                        # self.labels_logger.info(f"TOP LABEL     : city='{city}' | location='{location}' | neighborhood='{neighborhood}'")
                        # self.labels_logger.info(f"DOCUMENT MAP  : city='{cell_data.get('City')}' | location='{cell_data.get('Location')}' | neighborhood='{cell_data.get('Neighborhood')}'")
                        # self.labels_logger.info(f"{'='*80}\n")

                        # Check for mismatches between parsed label and document map
                        mismatch_msg = ""

                        if (
                            city
                            and city != cell_data.get("City")
                            or location
                            and location != cell_data.get("Location")
                            or neighborhood
                            and neighborhood != cell_data.get("Neighborhood")
                        ):
                            mismatch_msg += f"\n{'='*60}\n"
                            mismatch_msg += f"Mismatch for {self.sanitize_for_csv(self.strip_newlines_and_returns(txt1))}"

                        if city and city != cell_data.get("City"):
                            mismatch_msg += f"\n  City Mismatch:\n    Document Map City: '{cell_data.get('City')}'\n    Parsed Label City: '{city}'"
                        if location and location != cell_data.get("Location"):
                            mismatch_msg += f"\n  Location Mismatch:\n    Document Map Location: '{cell_data.get('Location')}'\n    Parsed Label Location: '{location}'"
                        if neighborhood and neighborhood != cell_data.get(
                            "Neighborhood"
                        ):
                            mismatch_msg += f"\n  Neighborhood Mismatch:\n    Document Map Neighborhood: '{cell_data.get('Neighborhood')}'\n    Parsed Label Neighborhood: '{neighborhood}'"

                        if mismatch_msg:
                            self.labels_logger.warning(mismatch_msg)

                    # Bottom Labels - Parse penny details
                    if txt2:
                        orientation, type_str, name_str = (
                            self.detect_orientation_and_type(txt2)
                        )

                        # Position is now correctly calculated in rows_from_vertical_pairs
                        # No need to adjust it here

                        # Create hash to detect duplicates (same location/name/orientation but different type)
                        stor_hash = f"{cell_data.get('Location')}|{self.sanitize_for_csv(self.strip_newlines_and_returns(name_str))}|{orientation}"

                        if stor_hash in position_stor:
                            self.logger.debug(
                                "Duplicate detected: Same Location/Name/Orientation with different Type (likely Copper vs Zinc Penny)"
                            )
                            position = position_stor.get(stor_hash)

                        row_to_append = cell_data.copy()
                        row_to_append.update({"Retired": "Yes" if retired else ""})
                        row_to_append.update(
                            {"Set #": current_set if current_set > 1 else ""}
                        )
                        row_to_append.update({"Orientation": orientation})
                        row_to_append.update({"Type": type_str})
                        row_to_append.update({"Position": position})
                        row_to_append.update(
                            {
                                "Name": self.sanitize_for_csv(
                                    self.strip_newlines_and_returns(name_str)
                                )
                            }
                        )

                        # Check if penny exists in database
                        is_new = not self.db.penny_exists(row_to_append)

                        # If penny is new, add it to database
                        if is_new:
                            self.db.add_penny(row_to_append)
                            self.logger.debug(
                                f"New penny found: {row_to_append['Name']} at {row_to_append['Location']}"
                            )
                        else:
                            self.logger.debug(
                                f"Existing penny: {row_to_append['Name']} at {row_to_append['Location']}"
                            )

                        # Add to output based on --new-only flag
                        if self.new_only:
                            # Only output new pennies
                            if is_new:
                                csv_rows.append(row_to_append)
                        else:
                            # Output all pennies
                            csv_rows.append(row_to_append)

                        position_stor.update({stor_hash: position})

        return csv_rows

    def write_csv(self, rows, out_path: str):
        """
        Write parsed rows into a CSV file.

        Args:
            rows (list of dict): Parsed data rows.
            out_path (str): Path for output CSV file.
        """
        header = [
            "State",
            "City",
            "Neighborhood",
            "Location",
            "Name",
            "Orientation",
            "Type",
            "Year",
            "Position",
            "Retired",
            "Set #",
            "Quantity",
            "Need",
        ]

        with open(out_path, self.write_mode, newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=header)
            # Only write header if creating new file
            if self.write_mode == "w":
                writer.writeheader()
            writer.writerows(rows)

    def run(
        self,
        input_file: str,
        output_file: str,
        append_loc: bool,
        new_only: bool = False,
        multi_line_dash: bool = False,
        write_mode_override: str = None,
    ):
        """
        Run the parser with input and output files.

        Args:
            input_file (str): Path to input DOCX file.
            output_file (str): Path to output CSV file.
            append_loc (bool): Whether to append location to neighborhood.
            new_only (bool): Only extract pennies not in database.
            multi_line_dash (bool): Allow dash separator in multi-line descriptions.
            write_mode_override (str): Override write mode ('w' or 'a'). If None, prompt user.
        """
        input_path = Path(input_file)
        if not input_path.exists():
            self.logger.error(f"Input file does not exist: {input_file}")
            return

        # Validate MIME type for Microsoft Word .docx format
        mime_type, _ = mimetypes.guess_type(input_file)
        if (
            mime_type
            != "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            # self.logger.error(f"Input file is not a valid Microsoft Word (.docx) file. Got MIME type: {mime_type}")
            print(
                f"ERROR: Input file is not a valid Microsoft Word (.docx) file. Got MIME type: {mime_type}"
            )
            return

        # Configure labels logger per input state/file (e.g., labels_dc.log)
        input_stem = Path(input_file).stem.lower()
        labels_log_file = f"labels_{input_stem}.log"
        self.labels_logger = self.setup_logging(
            labels_log_file, logger_name=f"{__name__}.labels", with_console=False
        )

        # Check if output file exists and get user preference before processing
        output_path = Path(output_file)
        write_mode = (
            write_mode_override or "w"
        )  # Default to write (overwrite) unless overridden

        if output_path.exists() and not write_mode_override:
            print(f"\nFile '{output_file}' already exists.")
            while True:
                choice = input(
                    "Choose an action:\n  1) Overwrite\n  2) Append\n  3) Backup and create new\n  4) Cancel\nEnter choice (1/2/3/4): "
                ).strip()

                if choice == "1":
                    if new_only:
                        print(
                            "\n⚠️  WARNING: Using 'Overwrite' with '--new-only' flag will DELETE all existing data!"
                        )
                        confirm = (
                            input("Are you sure you want to proceed? (yes/no): ")
                            .lower()
                            .strip()
                        )
                        if confirm == "no" or confirm == "n":
                            print("Operation cancelled.")
                            return
                    print("Overwriting existing file...")
                    write_mode = "w"
                    break
                elif choice == "2":
                    print("Appending to existing file...")
                    write_mode = "a"
                    break
                elif choice == "3":
                    # Create backup with timestamp
                    from datetime import datetime

                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    backup_path = (
                        output_path.parent
                        / f"{output_path.stem}_backup_{timestamp}{output_path.suffix}"
                    )
                    output_path.rename(backup_path)
                    print(f"Original file backed up to: {backup_path}")
                    write_mode = "w"
                    break
                elif choice == "4":
                    print("Operation cancelled.")
                    return
                else:
                    print("Invalid choice. Please enter 1, 2, 3, or 4.")

        self.short_location = append_loc
        self.write_mode = write_mode
        self.new_only = new_only
        self.multi_line_dash = multi_line_dash

        rows = self.parse_docx_to_rows(input_file)
        self.write_csv(rows, output_file)

        self.logger.info(f"Successfully parsed {len(rows)} rows.")
        self.logger.info(f"Output written to {output_file}")

    def run_file(
        self,
        input_file: str,
        output_file: str,
        append_loc: bool,
        multi_line_dash: bool,
        new_only: bool = False,
        write_mode_override: str = None,
    ):
        """Alias for run() method to match directory processing call signature."""
        return self.run(
            input_file,
            output_file,
            append_loc,
            new_only,
            multi_line_dash,
            write_mode_override,
        )


def parse_arguments(args=None):
    """
    Parse command-line arguments.

    Args:
        args: List of argument strings to parse. If None, uses sys.argv.

    Returns:
        Namespace object containing parsed arguments.
    """
    parser = argparse.ArgumentParser(
        description="Parse a pressed-penny DOCX file into a normalized CSV."
    )

    parser.add_argument(
        "--input",
        "-i",
        required=True,
        help="Path to input .docx file or directory containing .docx files",
    )

    parser.add_argument(
        "--output", "-o", required=True, help="Path to output .csv file"
    )

    parser.add_argument(
        "--short-location",
        "-sl",
        dest="short_loc",
        action="store_true",
        help="Keep Short Location if present. e.g. Big Top Toys instead of Buena Vista Street - Big Top Toys",
    )

    parser.add_argument(
        "--multi-line-dash",
        "-mld",
        dest="multi_line_dash",
        action="store_true",
        help="Will allow for dash separator in multi-line descriptions e.g. The Aristocats\n Something becomes The Aristocats - Something",
    )

    parser.add_argument(
        "--new-only",
        "-n",
        dest="new_only",
        action="store_true",
        help="Only extract pennies not already in the database",
    )

    return parser.parse_args(args)


def main():
    """
    Main entry point for the CLI.
    Uses argparse to collect command line arguments and perform parsing.
    """
    args = parse_arguments()

    penny_parser = PennyParser()

    # Handle directory or file input
    input_path = Path(args.input)

    if input_path.is_dir():
        # Process all .docx files in directory
        docx_files = sorted(input_path.glob("*.docx"))
        if not docx_files:
            print(f"No .docx files found in {args.input}")
            return

        for i, docx_file in enumerate(docx_files):
            print(f"\nProcessing {i+1}/{len(docx_files)}: {docx_file.name}")
            # For directory processing, always append (except first file)
            output_mode = "w" if i == 0 else "a"
            penny_parser.run_file(
                str(docx_file),
                args.output,
                args.short_loc,
                args.multi_line_dash,
                args.new_only,
                write_mode_override=output_mode,
            )
    else:
        # Process single file
        penny_parser.run(
            args.input,
            args.output,
            append_loc=args.short_loc,
            new_only=args.new_only,
            multi_line_dash=args.multi_line_dash,
        )


if __name__ == "__main__":
    main()
